# -*- coding: utf-8 -*-
"""生成软件著作权申请材料：
1. 源代码文档（完整 .docx + 前30页/后30页 .txt，每页50行）
2. 软件说明书（.txt + .docx）
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = Path(__file__).parent
OUTPUT_DIR = PROJECT_DIR / "软著材料"
OUTPUT_DIR.mkdir(exist_ok=True)

SOFTWARE_NAME = "基于地理位置的健康餐饮智能推荐系统"
VERSION = "V1.0"

code_files = [
    "index.html",
    "css/style.css",
    "js/config.js",
    "js/data/health-rules.js",
    "js/data/food-rules.js",
    "js/engine/scorer.js",
    "js/engine/storage.js",
    "js/engine/amap-service.js",
    "js/ui/render.js",
    "js/ui/app.js",
    "sw.js",
]

all_lines = []
for f in code_files:
    fp = PROJECT_DIR / f
    if not fp.exists():
        continue
    all_lines.append(f"// ===== 文件：{f} =====")
    with open(fp, "r", encoding="utf-8") as fh:
        for line in fh:
            s = line.rstrip()
            if s.strip() == "":
                continue
            all_lines.append(s)

print(f"有效代码行数: {len(all_lines)}")

PAGE = 50
front_lines = all_lines[: PAGE * 30]
back_lines = all_lines[-PAGE * 30:] if len(all_lines) > PAGE * 30 else all_lines


def write_txt(path, lines):
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"软件名称：{SOFTWARE_NAME} {VERSION}\n")
        f.write("=" * 60 + "\n\n")
        page = 1
        for i, line in enumerate(lines):
            f.write(line + "\n")
            if (i + 1) % PAGE == 0:
                f.write(f"\n{'-' * 30} 第 {page} 页 {'-' * 30}\n\n")
                page += 1


write_txt(OUTPUT_DIR / "源代码（前30页）.txt", front_lines)
write_txt(OUTPUT_DIR / "源代码（后30页）.txt", back_lines)


def build_code_docx(path, lines):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    title = doc.add_paragraph()
    run = title.add_run(f"{SOFTWARE_NAME} {VERSION}  源程序")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.save(path)


build_code_docx(OUTPUT_DIR / "源代码（完整）.docx", all_lines)
print(f"源代码文档已生成（共 {len(all_lines)} 行，前{len(front_lines)}/后{len(back_lines)}）")

manual = f"""软件说明书

软件名称：{SOFTWARE_NAME}
版本号：{VERSION}
开发完成日期：2026年6月

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

一、软件概述

1.1 软件用途
本软件是一款基于地理位置的健康餐饮智能推荐 Web 应用。系统结合高德地图定位与周边检索，从用户附近的真实餐饮店铺中，按"营养健康度"或"个人口味偏好"两种模式进行智能排序推荐，帮助用户在点外卖或就近就餐时做出更健康、更合口味的选择，并支持收藏与历史回溯。

1.2 开发环境
- 操作系统：Windows 10/11
- 编程语言：HTML5、CSS3、JavaScript（ES6）
- 架构方式：原生前端模块化（零运行时依赖）
- 第三方服务：高德地图开放平台 JS API 2.0（定位、地理编码、周边检索）
- 离线能力：PWA（Web App Manifest + Service Worker）
- 开发工具：Visual Studio Code

1.3 运行环境
- 操作系统：Windows / Linux / macOS / Android / iOS
- 现代浏览器（Chrome / Edge / Firefox / Safari）
- 网络连接（用于地图定位与店铺检索）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

二、主要功能

2.1 定位与地名搜索
系统支持两种定位方式：一是调用高德高精度定位获取用户当前位置（失败时回退到 IP 城市级定位）；二是通过地名输入框，将"中关村""哈工大"等地点经地理编码转换为坐标，查看指定地点附近的餐饮。

2.2 健康优选模式
从周边真实店铺中，按"健康度 + 距离"综合排序推荐，并提供均衡、减脂控碳、增肌高蛋白三种目标。系统依据店名与品类关键词为每家店计算"健康分"，对高油炸、高糖、高碳水的店铺标黄预警并给出搭配建议，对高蛋白/轻食类给予正向推荐。

2.3 周边美食模式
面向"想吃点什么"的场景，支持按关键词（火锅、日料、川菜、快餐等）、口味（不限/不吃辣/微辣/爱吃辣）、人数（按人数估算总价）、用餐时间（自动按当前钟点或手选早/午/晚/夜宵）多维筛选并排序。

2.4 真实评分与人均价格
系统读取高德地图的店铺详细信息，展示真实的用户评分、人均消费、营业时间、联系电话与店铺图片；在缺少真实数据时，按品类启发式估算人均价格区间作为兜底。

2.5 多维排序与评分筛选
结果支持"综合推荐 / 距离最近 / 评分最高 / 健康分最高"四种排序方式，并可按最低评分（如 4.0 分以上、4.5 分以上）筛选，帮助用户快速锁定优质店铺。

2.6 列表 / 地图双视图
用户可在"列表"与"地图"视图间一键切换。地图视图基于高德地图绘制用户位置与周边店铺标记，点击标记弹出店铺信息窗（名称、评分、人均、距离、地址），并自动调整视野完整展示所有结果。

2.7 分页加载更多
检索结果分页获取，用户可点击"加载更多附近店铺"持续追加结果（自动按店铺去重），突破单次检索条数限制。

2.8 收藏夹
用户可对任意店铺一键收藏，收藏数据本地持久化保存，并在独立的"我的收藏"栏目中集中查看与管理，可直接跳转高德查看店铺/导航。

2.9 搜索历史
系统自动记录最近的定位点与搜索地点，以标签形式展示，点击即可快速回到之前的位置重新检索，并对重复地点去重、超出上限自动截断。

2.10 平台跳转
提供"高德看店/导航"直达链接、电话一键拨打；并提供"美团外卖"辅助跳转（自动复制店名并打开外卖页面，便于粘贴搜索），移动端与桌面端分别适配。

2.11 离线使用（PWA）
通过 Web App Manifest 与 Service Worker 实现应用外壳的离线缓存，支持"添加到主屏幕"，弱网或离线时界面可秒开（地图与店铺数据仍需联网）。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

三、技术特点

3.1 分层模块化架构
项目划分为配置层、数据层（健康规则、美食规则）、引擎层（评分排序引擎、本地存储、地图服务封装）与界面层（渲染器、应用主入口），结构清晰、低耦合。

3.2 纯函数评分引擎与可配置权重
健康度评分、价格估算、两种模式的排序均以纯函数实现，并配套独立单元测试页自动验证；排序权重集中于配置文件统一管理，可直接调参而无需改动引擎代码。

3.3 Promise 化的地图服务封装
将高德 JS API 的脚本加载、定位、地理编码、周边检索统一封装为 Promise 接口，并对接口返回的错误信息进行透出，便于问题定位。

3.4 容错的本地存储
收藏夹与搜索历史基于 localStorage 持久化，并在浏览器存储不可用时降级为内存存储，不影响主流程。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

四、操作流程

4.1 打开应用，点击"定位并查找附近"获取当前位置，或在地名框输入地点后搜索。
4.2 在"健康优选/周边美食/我的收藏"三个视图间切换。
4.3 健康优选：选择均衡/减脂/增肌目标，查看按健康度+距离排序的推荐。
4.4 周边美食：设置想吃的品类、口味、人数、时间，点击"按条件找附近"。
4.5 通过结果控制条选择排序方式（综合/距离/评分/健康分）、设置最低评分筛选，或在"列表/地图"视图间切换。
4.6 调整搜索半径（1–5km）即时重新检索；点击"加载更多附近店铺"追加结果。
4.7 点击店铺卡片右上角星标收藏；在"我的收藏"中查看；点击历史标签回到旧位置。
4.8 点击"高德看店/导航""打电话"或"美团外卖"进行跳转。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

五、软件界面（截图位置，需手动插入）

[截图1：首页 - 定位与地名搜索]
[截图2：健康优选模式推荐列表]
[截图3：周边美食模式筛选与结果]
[截图4：店铺卡片（评分/人均/距离/电话）]
[截图5：地图视图 - 周边店铺标记]
[截图6：我的收藏]
[截图7：手机端界面]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

（说明书结束）
"""

with open(OUTPUT_DIR / "软件说明书.txt", "w", encoding="utf-8") as f:
    f.write(manual)


def build_manual_docx(path, text):
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"].font.size = Pt(11)
    for raw in text.split("\n"):
        line = raw.rstrip()
        p = doc.add_paragraph()
        if line.startswith("软件说明书"):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line and line[0] in "一二三四五六七八九十" and "、" in line[:3]:
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        else:
            p.add_run(line)
    doc.save(path)


build_manual_docx(OUTPUT_DIR / "软件说明书.docx", manual)
print("软件说明书已生成（txt + docx）")
print(f"\n所有材料保存在：{OUTPUT_DIR}")
